from typing import List, Dict
from dataclasses import dataclass

from openpyxl.workbook.workbook import Workbook
from openpyxl.worksheet.worksheet import Worksheet

from docx import Document

from exceptions import ReadWorksheetError

TC_SHEET_NAME = 'Sprint 1'
STORY_SHEET_NAME = 'story'
UAC_SHEET_NAME = 'uac'
TABLE_STYLE_NAME = 'Light Grid Accent 6' # https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
FIRST_ROW_TO_SCAN = 13

class TESTCASE_INDEX:
    ID = 0
    NAME = 1
    DESC = 3
    EXPECT_RESULT = 5
    ACTUAL_RESULT = 6

@dataclass
class Testcase:
    story_id: str
    name: str
    description: str
    expected_result: str
    actual_result: str

@dataclass
class Scenario:
    name: str
    tc_number: int

class TestFile:
    title :str
    story_id: str
    description: str
    _excel: Workbook

    uac = List[Dict]
    scenarios: List[Scenario]
    testcases: List[Testcase]

    def __init__(self, workbook: Workbook):
        self._excel = workbook
        pass

    def read_testfile(self, uac_sheet=True):
        workbook = self._excel
        try:
            story_data = self.get_story_data(workbook[STORY_SHEET_NAME])
            self.title = story_data['title']
            self.description = story_data['description']
            self.story_id = story_data['story_id']
            
            self.testcases = self.get_testcases(workbook[TC_SHEET_NAME])

            if uac_sheet:
                self.scenarios = self.get_scenarios_from_uac_sheet(workbook[UAC_SHEET_NAME])
            else:
                self.scenarios = self.get_scenarios_from_tc_sheet(workbook[TC_SHEET_NAME])
        except KeyError as err:
            if f'Worksheet {TC_SHEET_NAME} does not exist.' in str(err):
                raise ReadWorksheetError(TC_SHEET_NAME)

    def get_story_data(self, ws: Worksheet) -> Dict:
        story = {
            'title': '',
            'description': '',
            'story_id': ''
        }
        ws = ws[1:ws.max_row]
        for row in ws:
            story[row[0].value] = row[1].value
        
        return story

    def get_testcases(self, ws: Worksheet) -> List[Testcase]:
        testcases = []
        sliced_ws = ws[FIRST_ROW_TO_SCAN: ws.max_row-1]

        for index, row in enumerate(sliced_ws):
            if row[1].value is None:                # skip blank space between testcases.
                next_row = sliced_ws[index+1]   
                if next_row[0].value is None and next_row[1].value is None:       # last testcase followed with trail of blank spaces
                    break
                continue
            
            testcases.append(Testcase(
                story_id=row[TESTCASE_INDEX.ID].value,
                name=row[TESTCASE_INDEX.NAME].value,
                description=row[TESTCASE_INDEX.DESC].value,
                expected_result=row[TESTCASE_INDEX.EXPECT_RESULT].value,
                actual_result=row[TESTCASE_INDEX.ACTUAL_RESULT].value
            ))

        return testcases
    
    def get_scenarios_from_uac_sheet(self, ws: Worksheet) -> List[Scenario]:
        uacs = []
        for row in ws:
            uac, no = row[0].value, int(row[1].value)
            if uac is not None and no >= 0:
                uacs.append(Scenario(
                    name=uac,
                    tc_number=no
                ))
        return uacs

    def get_scenarios_from_tc_sheet(self, ws: Worksheet) -> List[Scenario]:
        uacs = []
        sliced_ws = ws[FIRST_ROW_TO_SCAN: ws.max_row-1]

        tc_number = 0
        for index, row in enumerate(sliced_ws):
            # check if its blank
            if row[0].value is None and row[1].value is None:
                continue

            # check if its scenario
            if row[0].value is not None and row[1].value is None:
                uacs.append(Scenario(name=row[0].value, tc_number=tc_number))
                continue

            # check if its tc
            tc_number += 1
        return uacs

    def write_document(self) -> Document:
        doc = Document()

        doc.add_heading(self.title, 0)
        story_id_paragraph = doc.add_paragraph('')
        story_id_paragraph.add_run(self.story_id).bold = True
        doc.add_paragraph(self.description)

        uacs = self.scenarios.copy()

        for index, testcase in enumerate(self.testcases):
            test_number = index + 1

            if len(uacs) > 0 and uacs[0].tc_number == test_number:
                doc.add_heading(uacs[0].name, 1)
                uacs.pop(0)
            
            doc.add_heading(f'Testcase{test_number}-{testcase.story_id}-{testcase.name}', 2)
            doc.add_paragraph(testcase.description)
            doc.add_paragraph("[ADD CONTENT HERE]")

            table = doc.add_table(rows=2, cols=2)
            table.style = TABLE_STYLE_NAME
            table.cell(0,0).text = 'Expected Results'
            table.cell(0,1).text = 'Actual Results'
            table.cell(1,0).text = testcase.expected_result or ""
            table.cell(1,1).text = testcase.actual_result or ""

        return doc

    