from typing import Dict, List

from openpyxl.workbook.workbook import Workbook
from openpyxl.worksheet.worksheet import Worksheet
from docx import Document
from dataclasses import dataclass

from exceptions import ReadWorksheetError

TC_SHEET_NAME = 'Sprint 1'
STORY_SHEET_NAME = 'story'
UAC_SHEET_NAME = 'uac'
TABLE_STYLE_NAME = 'Light Grid Accent 6' # https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
FIRST_ROW_TO_SCAN = 13

@dataclass
class Testcase:
    _id: str
    name: str
    description: str
    expected_result: str
    actual_result: str

    class INDEX:
        ID = 0
        NAME = 1
        DESC = 3
        EXPECT_RESULT = 5
        ACTUAL_RESULT = 6

class Scenario:
    name: str
    testcases: List[Testcase]

    def __init__(self, name = ''):
        self.name = name
        self.testcases = []
class TestFile:
    title :str
    story_id: str
    description: str
    scenario: List[Scenario]
    uac = List[Dict]
    _excel: Workbook
    
    class TESTCASE_INDEX:
        ID = 0
        NAME = 1
        DESC = 3
        EXPECT_RESULT = 5
        ACTUAL_RESULT = 6

    def __init__(self, workbook: Workbook = None, uac_sheet = False):
        self.scenario = []
        if workbook is not None:
            if type(workbook) is not Workbook:
                raise TypeError("testcase file is not Excel!")
            self._excel = workbook

            if uac_sheet:
                self.read_excel_tescase_v1()
            else:
                self.read_excel_tescase()

    def get_story_data(self, ws: Worksheet) -> None:
        story = {
            'title': '',
            'description': '',
            'story_id': ''
        }
        ws = ws[1:ws.max_row]
        for row in ws:
            story[row[0].value] = row[1].value

        self.title = story['title'] or ''
        self.story_id = story['story_id'] or ''
        self.description = story['description'] or ''

    def get_uac_from_sheet(self, ws:Worksheet) -> List[dict]:
        uacs = []
        for row in ws:
            uac, no = row[0].value, int(row[1].value)
            if uac is not None and no >= 0:
                uacs.append({
                    'uac': uac,
                    'no': no
                })

        self.uac = uacs
    
    def get_last_scenario(self):
        return self.scenario[-1]

    def read_testcases_v1(self, ws:Worksheet):
        uac = self.uac.copy()
        sliced_ws = ws[FIRST_ROW_TO_SCAN: ws.max_row-1]
        
        # read all testcase first
        self.scenario = []
        self.scenario.append(Scenario(name=uac[0]['uac']))
        uac.pop(0)

        testnumber = 0
        for index, row in enumerate(sliced_ws):
            if row[0].value is None:                # skip blank space between testcases.
                next_row = sliced_ws[index+1]   
                if next_row[0].value is None:       # last testcase followed with trail of blank spaces
                    break
                continue
            
            testnumber += 1
            print(testnumber)
            if len(uac) > 0 and testnumber == uac[0]['no']:
                print('[APPENDING]', uac[0]['no'])
                self.scenario.append(Scenario(name=uac[0]['uac']))
                uac.pop(0)

            testcase = Testcase(
                _id             = row[Testcase.INDEX.ID].value or '',
                name            = row[Testcase.INDEX.NAME].value or '',
                description     = row[Testcase.INDEX.DESC].value or '',
                expected_result = row[Testcase.INDEX.EXPECT_RESULT].value or '',
                actual_result   = row[Testcase.INDEX.ACTUAL_RESULT].value or ''
            )
            self.scenario[-1].testcases.append(testcase)

    def read_testcases(self, ws: Worksheet):
        def check_row_is_blank(row):
            if row[0].value is None and row[1].value is None:
                return True
            return False

        def check_row_is_uac(row):
            if isinstance(row[0].value, str) and row[1].value is None:
                return True
            return False

        curr_scenario = None
        testcases = []
        sliced_ws = ws[FIRST_ROW_TO_SCAN: ws.max_row-1]

        for index, row in enumerate(sliced_ws):
            if check_row_is_blank(row):
                next_row = sliced_ws[index+1]
                if check_row_is_blank(next_row):
                    curr_scenario.testcases = testcases
                    self.scenario.append(curr_scenario)
                    return
                continue
            
            if check_row_is_uac(row):
                if curr_scenario is None:
                    curr_scenario = Scenario()
                    curr_scenario.name = row[0].value or ''
                else:
                    curr_scenario.testcases = testcases
                    self.scenario.append(curr_scenario)
                    curr_scenario = Scenario()
                    curr_scenario.name = row[0].value or ''
                    testcases = []
                continue

            testcase = Testcase(
                _id             = row[Testcase.INDEX.ID].value or '',
                name            = row[Testcase.INDEX.NAME].value or '',
                description     = row[Testcase.INDEX.DESC].value or '',
                expected_result = row[Testcase.INDEX.EXPECT_RESULT].value or '',
                actual_result   = row[Testcase.INDEX.ACTUAL_RESULT].value or ''
            )
            testcases.append(testcase)

    def write_document(self) -> Document:
        doc = Document()
        
        doc.add_heading(self.title, 0)
        story_id_paragraph = doc.add_paragraph('')
        story_id_paragraph.add_run(self.story_id).bold = True
        doc.add_paragraph(self.description)

        index = 0
        for scenario in self.scenario:
            doc.add_heading(scenario.name, 1)

            for testcase in scenario.testcases:
                index += 1
                doc.add_heading(f'Testcase{index}-{self.story_id}-{testcase.name}', 2)
                doc.add_paragraph(testcase.description)
                doc.add_paragraph("[ADD CONTENT HERE]")

                table = doc.add_table(rows=2, cols=2)
                table.style = TABLE_STYLE_NAME
                table.cell(0,0).text = 'Expected Results'
                table.cell(0,1).text = 'Actual Results'
                table.cell(1,0).text = testcase.expected_result or ""
                table.cell(1,1).text = testcase.actual_result or ""
        
        return doc

    def read_excel_tescase(self):
        try:
            self.get_story_data(self._excel[STORY_SHEET_NAME])
            self.read_testcases(self._excel[TC_SHEET_NAME])
        except KeyError as err:
            if f'Worksheet {TC_SHEET_NAME} does not exist.' in str(err):
                raise ReadWorksheetError(TC_SHEET_NAME)

        if 'scenario' not in dir(self):
            raise ValueError('Scenario not found!')

    def read_excel_tescase_v1(self):
        try:
            self.get_story_data(self._excel[STORY_SHEET_NAME])
            self.get_uac_from_sheet(self._excel[UAC_SHEET_NAME])
            self.read_testcases_v1(self._excel[TC_SHEET_NAME])
        except KeyError as err:
            if f'Worksheet {TC_SHEET_NAME} does not exist.' in str(err):
                raise ReadWorksheetError(TC_SHEET_NAME)

        if 'scenario' not in dir(self):
            raise ValueError('Scenario not found!')

    def get_docx(self) -> Document:
        return self.write_document()