#By: Nabeel Kahlil Maulana

import sys
import os
from typing import List

from openpyxl import load_workbook
from docx import Document
from openpyxl.worksheet.worksheet import Worksheet
from openpyxl.workbook.workbook import Workbook

TC_SHEET_NAME = 'Sprint 1'
STORY_SHEET_NAME = 'story'
UAC_SHEET_NAME = 'uac'
TABLE_STYLE_NAME = 'Light Grid Accent 6' # https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
FIRST_ROW_TO_SCAN = 13

class ReadWorksheetError(Exception):
    """Cannot find Testcase worksheet"""
    def __init__(self, worksheet_name = None):
        self.worksheet_name = worksheet_name
        tc_name = f'"{self.worksheet_name}"' if worksheet_name else ''
        self.message = f"Cannot find Testcase {tc_name} Worksheet"

class TESTCASE_INDEX:
    ID = 0
    NAME = 1
    DESC = 3
    EXPECT_RESULT = 5
    ACTUAL_RESULT = 6

def read_testcases(ws: Worksheet) -> dict:
    testcases = []
    sliced_ws = ws[FIRST_ROW_TO_SCAN: ws.max_row-1]

    for index, row in enumerate(sliced_ws):
        if row[0].value is None:                # skip blank space between testcases.
            next_row = sliced_ws[index+1]   
            if next_row[0].value is None:       # last testcase followed with trail of blank spaces
                break
            continue

        testcases.append({
            'ID'                : row[TESTCASE_INDEX.ID].value,
            'name'              : row[TESTCASE_INDEX.NAME].value,
            'description'       : row[TESTCASE_INDEX.DESC].value,
            'expected_result'   : row[TESTCASE_INDEX.EXPECT_RESULT].value,
            'actual_result'     : row[TESTCASE_INDEX.ACTUAL_RESULT].value
        })

    return testcases

def get_story_data(ws: Worksheet) -> dict:
    story = {
        'title': '',
        'description': '',
        'story_id': ''
    }
    ws = ws[1:ws.max_row]
    for row in ws:
        story[row[0].value] = row[1].value
    return story

def get_uac(ws: Worksheet) -> list:
    uacs = []
    for row in ws:
        uac, no = row[0].value, int(row[1].value)
        if uac is not None and no >= 0:
            uacs.append({
                'uac': uac,
                'no': no
            })
    return uacs

def write_document(story: dict, testcases: list, uacs: List[dict]) -> Document:
    doc = Document()
    
    doc.add_heading(story['title'], 0)
    story_id_paragraph = doc.add_paragraph('')
    story_id_paragraph.add_run(story['story_id']).bold = True
    doc.add_paragraph(story['description'])

    for index, testcase in enumerate(testcases):
        test_number = index + 1

        if len(uacs) > 0 and uacs[0]['no'] == test_number:
            doc.add_heading(uacs[0]['uac'], 1)
            uacs.pop(0)
        
        doc.add_heading(f'Testcase{test_number}-{story["story_id"]}-{testcase["name"]}', 2)
        doc.add_paragraph(testcase['description'])
        doc.add_paragraph("[ADD CONTENT HERE]")

        table = doc.add_table(rows=2, cols=2)
        table.style = TABLE_STYLE_NAME
        table.cell(0,0).text = 'Expected Results'
        table.cell(0,1).text = 'Actual Results'
        table.cell(1,0).text = testcase['expected_result'] or ""
        table.cell(1,1).text = testcase['actual_result'] or ""
    
    return doc

def windows_to_unix_path(path: str):
    if '/' in path:
        return path
    return path.replace('\\', '/')

def convert(workbook: Workbook):
    try:
        story = get_story_data(workbook[STORY_SHEET_NAME])
        testcases = read_testcases(workbook[TC_SHEET_NAME])
        uacs = get_uac(workbook[UAC_SHEET_NAME])
    except KeyError as err:
        if f'Worksheet {TC_SHEET_NAME} does not exist.' in str(err):
            raise ReadWorksheetError(TC_SHEET_NAME)

    doc = write_document(story, testcases, uacs)
    return doc

def rename_tc_filename(filename: str):
    filename = list(filename)
    filename[0:2] = 'SS'
    return ''.join(filename)

def save_document(doc, result_location):
    if 'results' not in os.listdir():
        os.makedirs(os.join(os.getcwd(),'results'))
    doc.save(result_location)

def command_line_app(source_location, target_directory):
    file_name = os.path.basename(source_location)
    target_filename = file_name.split('.')[0] + '.docx'

    if 'TC' == target_filename[0:2]:
        target_filename = rename_tc_filename(target_filename)
    
    workbook = load_workbook(filename=source_location, read_only=True)
    doc = convert(workbook)

    target_filepath = os.path.join(target_directory, target_filename)
    save_document(doc, target_filepath)

    print(f'file saved in {target_filepath}')

if __name__ == '__main__':
    argv = sys.argv
    try:
        source_location = argv[1]
        target_directory = argv[2]
    except IndexError:
        if len(argv) < 2:
            raise Exception("argument required: target source is not defined")
        if len(argv) < 3:
            target_directory = os.path.join(os.getcwd(), 'results')
    
    source_location = windows_to_unix_path(source_location)
    target_directory = windows_to_unix_path(target_directory)

    command_line_app(source_location, target_directory)
